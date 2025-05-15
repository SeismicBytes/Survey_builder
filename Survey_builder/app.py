# app.py (Updated for .docx and file uploads)
import os
import sys
import logging
from flask import Flask, request, jsonify
from flask_cors import CORS
from dotenv import load_dotenv
from io import BytesIO # Needed for reading file stream in memory
from docx import Document # Needed for reading .docx files

import spg_parser # Import your (placeholder) parser script
import ai_client # Your parser needs this

# --- Load Environment Variables ---
load_dotenv()
print("Attempted to load .env file.")

# --- Configuration ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(name)s - %(message)s')
logger = logging.getLogger(__name__)

# --- AI Client Configuration (CRITICAL!) ---
API_KEY = os.getenv("GEMINI_API_KEY")
if not API_KEY:
    logger.critical("FATAL ERROR: GEMINI_API_KEY not found in environment variables or .env file.")
    raise EnvironmentError("GEMINI_API_KEY not found. Ensure it's set in the environment or a .env file.")
else:
     masked_key = API_KEY[:4] + "****" + API_KEY[-4:]
     logger.info(f"GEMINI_API_KEY found (masked: {masked_key}). Attempting AI client configuration...")

try:
    # Replace with your actual configuration call
    ai_client.configure(api_key=API_KEY)
    logger.info("AI Client configuration function called.")
    if not hasattr(ai_client, 'is_client_configured') or not ai_client.is_client_configured():
         logger.warning("ai_client.is_client_configured() check failed or returned false after configuration attempt.")
    else:
        logger.info("AI Client appears to be configured successfully.")
except AttributeError as ae:
     logger.critical(f"Configuration Error: Missing function in ai_client - {ae}", exc_info=True)
     raise RuntimeError(f"AI Client configuration failed: Missing function in ai_client - {ae}")
except Exception as config_err:
    logger.critical(f"Failed to configure AI Client during setup: {config_err}", exc_info=True)
    raise RuntimeError(f"Failed to configure AI Client: {config_err}")

# --- Helper Function for Text Extraction ---
def extract_text_with_markdown_tables(docx_file_stream):
    """
    Reads a .docx file stream and extracts text, converting tables to Markdown format.
    Returns the extracted text as a single string or None on failure.
    """
    try:
        document = Document(docx_file_stream)
        full_text_parts = []
        logger.info("Processing .docx file content...")

        # Helper to process table and append as markdown
        def process_table(table):
            markdown_table = []
            try:
                # Escape pipe characters in header and determine column count
                header_cells = [cell.text.strip().replace('|', '\\|') for cell in table.rows[0].cells]
                col_count = len(header_cells)
                if col_count == 0: return # Skip empty tables

                markdown_table.append("| " + " | ".join(header_cells) + " |")
                markdown_table.append("| " + " | ".join(['---'] * col_count) + " |")

                for row in table.rows[1:]:
                    # Escape pipe characters in row cells
                    row_cells = [cell.text.strip().replace('|', '\\|') for cell in row.cells]
                    # Pad or truncate row cells to match header column count
                    current_row_col_count = len(row_cells)
                    if current_row_col_count < col_count:
                        row_cells.extend([''] * (col_count - current_row_col_count))
                    elif current_row_col_count > col_count:
                        row_cells = row_cells[:col_count]

                    markdown_table.append("| " + " | ".join(row_cells) + " |")

                # Add explicit markers for AI
                return "\n\n**DOCX_TABLE_START**\n" + "\n".join(markdown_table) + "\n**DOCX_TABLE_END**\n\n"
            except Exception as table_err:
                logger.error(f"Error processing a table: {table_err}", exc_info=True)
                return "\n\n**DOCX_TABLE_ERROR**\n[Error processing table content]\n**DOCX_TABLE_ERROR_END**\n\n" # Indicate table error

        # Iterate through elements preserving approximate order (paragraphs and tables)
        # This approach iterates through the underlying XML structure
        for element in document.element.body:
            if element.tag.endswith('tbl'):
                 # Find the corresponding Table object based on the element
                 table_object = None
                 for t in document.tables:
                    if t._element is element:
                         table_object = t
                         break
                 if table_object:
                    table_md = process_table(table_object)
                    if table_md:
                         full_text_parts.append(table_md)
            elif element.tag.endswith('p'):
                # Find the corresponding Paragraph object based on the element
                para_object = None
                for p in document.paragraphs:
                     if p._element is element:
                         para_object = p
                         break
                if para_object and para_object.text.strip(): # Add non-empty paragraphs
                     full_text_parts.append(para_object.text)

        logger.info(f"Finished processing .docx. Extracted {len(full_text_parts)} parts.")
        # Join parts with double newline for better separation, especially around tables
        return "\n\n".join(full_text_parts)

    except Exception as e:
        logger.error(f"Error reading or processing docx file stream: {e}", exc_info=True)
        return None # Indicate failure

# --- Flask App Setup ---
app = Flask(__name__)
CORS(app)

# --- API Endpoint (Modified for File Upload) ---
@app.route('/parse', methods=['POST'])
def parse_survey_file():
    """
    API endpoint to parse uploaded survey files (.txt or .docx).
    Expects FormData with a 'file' part and optional 'run_id'.
    Returns JSON output from spg_parser.parse_spg or an error message.
    """
    run_id = request.form.get('run_id', 'API_FILE_RUN') # Get run_id from form data
    logger.info(f"[{run_id}] Received request on /parse endpoint (expecting file) from {request.remote_addr}")

    # --- AI Client Check ---
    if not ai_client.is_client_configured():
        logger.error(f"[{run_id}] Aborting request: AI Client is not configured.")
        return jsonify({"error": "Server configuration error: AI Client not ready."}), 503

    # --- File Handling ---
    if 'file' not in request.files:
        logger.error(f"[{run_id}] No 'file' part in the request.")
        return jsonify({"error": "No file part found in the request"}), 400

    file = request.files['file']

    if file.filename == '':
        logger.error(f"[{run_id}] No file selected for upload.")
        return jsonify({"error": "No file selected"}), 400

    logger.info(f"[{run_id}] Received file: {file.filename}")
    extracted_text = None
    file_extension = ""

    try:
        filename_lower = file.filename.lower()
        if filename_lower.endswith('.docx'):
            file_extension = '.docx'
            logger.info(f"[{run_id}] Processing .docx file...")
            # Read file content as bytes and wrap in BytesIO for python-docx
            file_stream = BytesIO(file.read())
            extracted_text = extract_text_with_markdown_tables(file_stream)
            if extracted_text is None: # Check if extraction failed
                 logger.error(f"[{run_id}] Failed to extract text from .docx file.")
                 return jsonify({"error": f"Failed to extract text from DOCX file: {file.filename}"}), 500
        elif filename_lower.endswith('.txt'):
            file_extension = '.txt'
            logger.info(f"[{run_id}] Processing .txt file...")
            # Read text file, decode assuming UTF-8
            extracted_text = file.read().decode('utf-8')
        else:
            logger.error(f"[{run_id}] Unsupported file type: {file.filename}")
            return jsonify({"error": "Unsupported file type. Please upload .txt or .docx"}), 415 # Unsupported Media Type

        logger.info(f"[{run_id}] Successfully extracted text from {file_extension} (length: {len(extracted_text)} characters).")

    except UnicodeDecodeError:
         logger.error(f"[{run_id}] Error decoding .txt file. Ensure it's UTF-8 encoded.")
         return jsonify({"error": "Failed to decode .txt file. Please ensure it is UTF-8 encoded."}), 400
    except Exception as file_err:
         logger.error(f"[{run_id}] Error reading file {file.filename}: {file_err}", exc_info=True)
         return jsonify({"error": f"Error reading file: {file_err}"}), 500


    # --- Parsing ---
    if extracted_text is None: # Should have been caught above, but double-check
         logger.error(f"[{run_id}] Extracted text is None after file processing. Cannot parse.")
         return jsonify({"error": "Internal server error: Failed to get text content."}), 500

    try:
        batch_size = 15000 # Keep batch size reasonable, adjust as needed
        logger.info(f"[{run_id}] Calling spg_parser.parse_spg with batch size {batch_size}...")

        # Call your main parsing function (which uses the mock AI client for now)
        result_json = spg_parser.parse_spg(
            full_text=extracted_text,
            batch_char_size=batch_size,
            run_id=run_id
        )

        if result_json is None:
            logger.error(f"[{run_id}] spg_parser.parse_spg returned None. Check parser logs.")
            return jsonify({"error": "Parsing failed within the parser module. Check server logs."}), 500

        logger.info(f"[{run_id}] Parsing call completed. Returning result JSON.")
        return jsonify(result_json), 200

    except Exception as parse_err:
        logger.error(f"[{run_id}] Exception during parsing call: {parse_err}", exc_info=True)
        return jsonify({"error": f"An unexpected error occurred on the server during parsing: {parse_err}"}), 500

# --- Run the App ---
if __name__ == '__main__':
    logger.info("Starting Flask server...")
    flask_debug = os.getenv('FLASK_DEBUG', 'False').lower() in ('true', '1', 't')
    app.run(host='0.0.0.0', port=5000, debug=flask_debug)